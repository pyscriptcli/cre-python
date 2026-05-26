# ==========================================
# DELTA | PREMIUM DUE DILIGENCE PIPELINE
# ==========================================

import streamlit as st
import io
import re
import math
import hashlib
from datetime import datetime
import numpy as np

# Core Geospatial, Document, & Vision Engine Binaries
import pymupdf as fitz  # FIXED: Swapped to explicit PyMuPDF binding layer
from docx import Document
from docx.shared import Pt
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter
import matplotlib.pyplot as plt
import shutil
import os

# Explicit dynamic resolution of Linux system path bindings for Tesseract
if not shutil.which("tesseract"):
    # Fallback default hardcoded paths for standard Debian/Ubuntu containers
    fallback_paths = ["/usr/bin/tesseract", "/usr/local/bin/tesseract"]
    for path in fallback_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break
else:
    pytesseract.pytesseract.tesseract_cmd = shutil.which("tesseract")

# ==========================================
# BLOCK 1: STATE HYDRATION
# ==========================================
if 'initialized' not in st.session_state:
    st.session_state.initialized = True
    st.session_state.uploaded_titles = {}
    st.session_state.title_hashes = {}
    st.session_state.title_order = []           
    st.session_state.title_roles = {}           
    st.session_state.dd_processing_complete = False 
    st.session_state.base_title_idx = 0      
    st.session_state.counter_title_idx = 1       
    st.session_state.extracted_tech_desc = ""
    st.session_state.parsed_polygon = None

# ==========================================
# BLOCK 2: SYSTEM DESIGN SYSTEM (CSS)
# ==========================================
def inject_luxury_system_css():
    st.markdown("""
        <style>
            @import url('https://fonts.googleapis.com/css2?family=Cinzel:wght=400;500&family=Inter:wght@300;400;500;600&display=swap');
            
            .stApp {
                background-color: #F4F5F7 !important;
                color: #1A1A1A !important;
                font-family: 'Inter', sans-serif !important;
            }
            
            [data-testid="stHeader"] { background-color: rgba(0,0,0,0) !important; border-bottom: none !important; }
            footer {visibility: hidden !important;}
            .block-container { padding-top: 2rem !important; padding-bottom: 3rem !important; max-width: 95% !important; }
            
            [data-testid="stSidebar"] {
                background-color: #FFFFFF !important;
                border-right: 1px solid #E0E0E0;
            }
            
            .brand-title {
                font-family: 'Cinzel', serif !important;
                color: #1A1A1A !important;
                letter-spacing: 4px;
                font-weight: 500;
                font-size: 28px;
                margin-bottom: 0.2rem;
            }
            .brand-subtitle { font-family: 'Inter', sans-serif; color: #737373; font-size: 10px; letter-spacing: 2px; text-transform: uppercase; margin-bottom: 2.5rem; }
            
            .doc-cell {
                background-color: #FFFFFF;
                padding: 24px 32px;
                border: 1px solid #E0E0E0;
                box-shadow: 0 4px 8px rgba(0,0,0,0.02);
                border-radius: 6px;
                height: 100%;
                margin-bottom: 16px;
            }
            .empty-cell { background-color: transparent !important; border: none !important; box-shadow: none !important; }
            
            .landing-wrapper {
                display: flex; flex-direction: column; align-items: center; justify-content: center; text-align: center;
                width: 100%; margin: 0 auto; padding: 4rem 2rem; background-color: #FFFFFF;
                border: 1px solid #E0E0E0; box-shadow: 0 4px 12px rgba(0, 0, 0, 0.05);
            }
            
            .stFileUploader { border: 1px dashed #1A1A1A !important; background-color: #F4F5F7 !important; padding: 1.5rem !important; }
            .stButton > button {
                background-color: transparent !important; color: #1A1A1A !important; border: 1px solid #1A1A1A !important;
                border-radius: 0px !important; font-family: 'Inter', sans-serif; font-size: 11px !important; letter-spacing: 1.5px;
                text-transform: uppercase; padding: 0.6rem 2rem !important; transition: all 0.3s ease; width: 100%;
            }
            .stButton > button:hover { background-color: #1A1A1A !important; color: #FFFFFF !important; }
            
            .stream-paragraph { color: #1A1A1A; font-size: 13px; line-height: 1.8; word-wrap: break-word; margin-bottom: 0px; }
            .add-token { background-color: #D1FAE5 !important; color: #065F46 !important; padding: 2px 4px; border-radius: 2px; }
            .del-token { background-color: #FEE2E2 !important; color: #991B1B !important; text-decoration: line-through; padding: 2px 4px; border-radius: 2px; }
            .trace-flag { font-family: 'Cinzel', serif; color: #737373; font-size: 10px; letter-spacing: 1px; text-transform: uppercase; margin-bottom: 0.8rem; display: block; border-bottom: 1px solid #F4F5F7; padding-bottom: 6px; }
            
            .advisory-panel { background-color: #FFFFFF; border: 1px solid #E0E0E0; box-shadow: 0 2px 6px rgba(0,0,0,0.02); padding: 1rem; margin-bottom: 1rem; border-radius: 6px; }
            .advisory-header { font-family: 'Inter', sans-serif; font-weight: 600; color: #1A1A1A; font-size: 12px; margin-bottom: 0.5rem; }
            
            .minimap-row { width: 100%; height: 8px; border-radius: 1px; margin-top: 30px; }
            .minimap-equal { background-color: #EAECEF; opacity: 0.4; }
            .minimap-replace { background-color: #FCA5A5; }
            .minimap-delete { background-color: #EF4444; }
            .minimap-insert { background-color: #10B981; }
            
            .crypto-banner { font-family: monospace; font-size: 10px; background-color: #FFFFFF; border: 1px solid #E0E0E0; padding: 8px 12px; color: #737373; margin-bottom: 1.5rem; border-radius: 4px; }
            div[data-baseweb="select"] { background-color: #FFFFFF !important; border-radius: 0px !important; }
            input { border-radius: 4px !important; background-color: #FFFFFF !important; color: #1A1A1A !important; border: 1px solid #E0E0E0 !important; font-size: 12px !important; padding: 8px 10px !important; }
            textarea { border-radius: 4px !important; background-color: #FFFFFF !important; color: #1A1A1A !important; border: 1px solid #E0E0E0 !important; font-size: 12px !important; }
        </style>
    """, unsafe_allow_html=True)

# ==========================================
# BLOCK 3: ENGINE PARSING, OCR & HASHING
# ==========================================
def preprocess_image_for_ocr(image_bytes):
    image = Image.open(io.BytesIO(image_bytes)).convert('L')
    image = image.filter(ImageFilter.SHARPEN)
    enhancer = ImageEnhance.Contrast(image)
    image = enhancer.enhance(2.0)
    return image

def extract_text_via_ocr(file_bytes):
    try:
        img = preprocess_image_for_ocr(file_bytes)
        text = pytesseract.image_to_string(img)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        return lines
    except Exception as e:
        return [f"[OCR Error Engine Exception: {str(e)}]"]

def parse_pdf(file_bytes):
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    paragraphs = []
    for page in doc:
        text_blocks = page.get_text("blocks")
        if len(text_blocks) > 0:
            for b in text_blocks:
                text = b[4].strip()
                if text: paragraphs.append(" ".join(text.split()))
        else:
            pix = page.get_pixmap(dpi=200)
            img_data = pix.tobytes("png")
            paragraphs.extend(extract_text_via_ocr(img_data))
    return paragraphs

def parse_docx(file_bytes):
    doc = Document(io.BytesIO(file_bytes))
    paragraphs = []
    for p in doc.paragraphs:
        if p.text.strip(): paragraphs.append(" ".join(p.text.split()))
    return paragraphs

def load_due_diligence_matrices(uploaded_files):
    for file in uploaded_files:
        if file.name not in st.session_state.uploaded_titles:
            bytes_data = file.read()
            file_hash = hashlib.sha256(bytes_data).hexdigest()
            st.session_state.title_hashes[file.name] = file_hash
            
            if file.name.endswith('.pdf'):
                parsed_text = parse_pdf(bytes_data)
            elif file.name.endswith('.docx'):
                parsed_text = parse_docx(bytes_data)
            elif file.name.lower().endswith(('.png', '.jpg', '.jpeg', '.tiff', '.bmp')):
                parsed_text = extract_text_via_ocr(bytes_data)
            else:
                continue
            
            st.session_state.uploaded_titles[file.name] = parsed_text
            if file.name not in st.session_state.title_order:
                st.session_state.title_order.append(file.name)
                st.session_state.title_roles[file.name] = "Certified True Copy (Base)"

# ==========================================
# BLOCK 4: FUZZY ALIGNMENT ENGINE
# ==========================================
def compute_fuzzy_alignment_matrix(left_paras, right_paras, threshold=0.35):
    import difflib
    matched_right_indices = set()
    alignment_opcodes = []
    
    for i, lp in enumerate(left_paras):
        best_ratio = 0.0
        best_j = None
        lp_tokens = sorted(lp.lower().split())
        
        for j, rp in enumerate(right_paras):
            if j in matched_right_indices: continue
            rp_tokens = sorted(rp.lower().split())
            ratio = difflib.SequenceMatcher(None, lp_tokens, rp_tokens).ratio()
            if ratio > best_ratio:
                best_ratio = ratio
                best_j = j
                
        if best_ratio >= threshold and best_j is not None:
            matched_right_indices.add(best_j)
            if best_ratio > 0.95 and lp == right_paras[best_j]: 
                alignment_opcodes.append(('equal', i, i, best_j, best_j))
            else: 
                alignment_opcodes.append(('replace', i, i, best_j, best_j))
        else:
            alignment_opcodes.append(('delete', i, i, None, None))
            
    for j in range(len(right_paras)):
        if j not in matched_right_indices:
            alignment_opcodes.append(('insert', None, None, j, j))
            
    alignment_opcodes.sort(key=lambda x: (x[1] if x[1] is not None else float('inf'), x[3] if x[3] is not None else 0))
    return alignment_opcodes

def compute_token_diff_html(text1, text2):
    import difflib
    matcher = difflib.SequenceMatcher(None, text1.split(), text2.split())
    out1, out2 = [], []
    for tag, i1, i2, j1, j2 in matcher.get_opcodes():
        w1 = " ".join(text1.split()[i1:i2]); w2 = " ".join(text2.split()[j1:j2])
        if tag == 'equal': out1.append(w1); out2.append(w2)
        elif tag == 'replace': 
            out1.append(f'<span class="del-token">{w1}</span>')
            out2.append(f'<span class="add-token">{w2}</span>')
        elif tag == 'delete': out1.append(f'<span class="del-token">{w1}</span>')
        elif tag == 'insert': out2.append(f'<span class="add-token">{w2}</span>')
    return " ".join(out1), " ".join(out2)

# ==========================================
# BLOCK 5: TECHNICAL DESCRIPTION GEOMETRY ENGINE
# ==========================================
def parse_technical_description(text_content):
    cleaned = text_content.replace('\n', ' ')
    pattern = r'THENCE\s+([N|S])\s*[\.]?\s*(\d+)\s*DEG\.\s*(\d+)\'?\s*([E|W])\s*[\.]?[\, ]?\s*(\d+(?:\.\d+)?)\s*M\.'
    matches = re.findall(pattern, cleaned, re.IGNORECASE)
    
    if not matches:
        pattern_fallback = r'(N|S)\s*(\d+)\s*DEG\s*(\d+)\s*(E|W)\s*(\d+(?:\.\d+)?)'
        matches = re.findall(pattern_fallback, cleaned, re.IGNORECASE)
    
    current_x, current_y = 0.0, 0.0
    coordinates = [(current_x, current_y)]
    
    for i, (ns, deg, mins, ew, dist) in enumerate(matches):
        try:
            d_val = float(dist)
            deg_val = float(deg)
            min_val = float(mins)
            
            angle_rad = math.radians(deg_val + (min_val / 60.0))
            ns_sign = 1.0 if ns.upper() == 'N' else -1.0
            ew_sign = 1.0 if ew.upper() == 'E' else -1.0
            
            dy = d_val * math.cos(angle_rad) * ns_sign
            dx = d_val * math.sin(angle_rad) * ew_sign
            
            current_x += dx
            current_y += dy
            coordinates.append((current_x, current_y))
        except ValueError:
            continue
            
    if len(coordinates) > 1:
        coordinates[-1] = (0.0, 0.0)
        
    return coordinates

def generate_white_blueprint_pdf(coordinates):
    doc = fitz.open()
    page_w, page_h = 612, 792 
    page = doc.new_page(width=page_w, height=page_h)
    
    page.draw_rect(fitz.Rect(0, 0, page_w, page_h), color=None, fill=(1, 1, 1))
    page.insert_text(fitz.Point(54, 54), "DELTA AUTOMATED LOT SURVEY PLAN", fontsize=14, fontname="Helvetica-Bold", color=(0.1, 0.1, 0.1))
    page.insert_text(fitz.Point(54, 70), f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Architectural Matrix Pipeline", fontsize=8, fontname="Helvetica-Oblique", color=(0.4, 0.4, 0.4))
    
    if len(coordinates) > 1:
        x_vals = [c[0] for c in coordinates]
        y_vals = [c[1] for c in coordinates]
        
        min_x, max_x = min(x_vals), max(x_vals)
        min_y, max_y = min(y_vals), max(y_vals)
        
        span_x = max_x - min_x if max_x != min_x else 1.0
        span_y = max_y - min_y if max_y != min_y else 1.0
        
        scale = min(400.0 / span_x, 400.0 / span_y)
        center_x, center_y = 306, 430
        
        map_cx = min_x + (span_x / 2.0)
        map_cy = min_y + (span_y / 2.0)
        
        transformed_pts = []
        for x, y in coordinates:
            px = center_x + (x - map_cx) * scale
            py = center_y - (y - map_cy) * scale 
            transformed_pts.append(fitz.Point(px, py))
            
        page.draw_line(fitz.Point(center_x-20, center_y), fitz.Point(center_x+20, center_y), color=(0.8, 0.8, 0.8), width=0.5)
        page.draw_line(fitz.Point(center_x, center_y-20), fitz.Point(center_x, center_y+20), color=(0.8, 0.8, 0.8), width=0.5)
        
        for idx in range(len(transformed_pts) - 1):
            p1 = transformed_pts[idx]
            p2 = transformed_pts[idx+1]
            page.draw_line(p1, p2, color=(0.1, 0.1, 0.1), width=1.5)
            page.draw_circle(p1, 2.5, color=(0.1, 0.1, 0.1), fill=(1, 1, 1))
            page.insert_text(fitz.Point(p1.x + 4, p1.y - 4), f"P{idx+1}", fontsize=7, color=(0.2, 0.2, 0.2))
            
    page.draw_line(fitz.Point(54, 720), fitz.Point(558, 720), color=(0.9, 0.9, 0.9), width=1)
    page.insert_text(fitz.Point(54, 740), "Disclaimer: This plot is generated directly from string matching parameters from document technical transcripts for reference use.", fontsize=7, color=(0.5, 0.5, 0.5))
    
    return doc.write()

def generate_kml_payload(coordinates):
    base_lat, base_lng = 14.6541, 120.9791 
    
    kml_str = """<?xml version="1.0" encoding="UTF-8"?>
<kml xmlns="http://www.opengis.net/kml/2.2">
  <Document>
    <name>DELTA Due Diligence Boundary Plan Extracted Lot</name>
    <Style id="lotLineStyle">
      <LineStyle><color>ff0000ff</color><width>3</width></LineStyle>
      <PolyStyle><color>40ffffff</color></PolyStyle>
    </Style>
    <Placemark>
      <name>Extracted Boundary Poly</name>
      <styleUrl>#lotLineStyle</styleUrl>
      <Polygon>
        <outerBoundaryIs>
          <LinearRing>
            <coordinates>
"""
    for x, y in coordinates:
      lng_offset = x / 111320.0
      lat_offset = y / 111054.0
      kml_str += f"              {base_lng + lng_offset},{base_lat + lat_offset},0\n"
        
    kml_str += """            </coordinates>
          </LinearRing>
        </outerBoundaryIs>
      </Polygon>
    </Placemark>
  </Document>
</kml>"""
    return kml_str.encode('utf-8')

# ==========================================
# BLOCK 6: HIGH-FIDELITY EXPORTERS
# ==========================================
def export_due_diligence_docx(left_paras, right_paras, title_left, title_right, alignment_opcodes):
    doc = Document()
    section = doc.sections[-1]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    
    doc.add_paragraph().add_run("DELTA DUE DILIGENCE RISK ANALYSIS TREE").bold = True
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = f"REFERENCE CORE ({title_left})"
    hdr_cells[1].text = f"COUNTER EVAL VALUE ({title_right})"
    hdr_cells[2].text = "VARIANCE ANALYSIS / RED FLAG MATRIX"
    
    for tag, i1, _, j1, _ in alignment_opcodes:
        row = table.add_row()
        if tag == 'equal':
            row.cells[0].text = left_paras[i1] if i1 is not None else ""
            row.cells[1].text = right_paras[j1] if j1 is not None else ""
            row.cells[2].text = "Verified Match. Zero variance found."
        elif tag == 'replace':
            row.cells[0].text = left_paras[i1] if i1 is not None else ""
            row.cells[1].text = right_paras[j1] if j1 is not None else ""
            row.cells[2].text = "Mismatched Parameter Text block. Variance verified requires audit confirmation."
        elif tag == 'delete':
            row.cells[0].text = left_paras[i1] if i1 is not None else ""
            row.cells[1].text = "[Omitted Field Frame]"
            row.cells[2].text = "Red Flag Structural Data point absent from target evaluation document."
        elif tag == 'insert':
            row.cells[0].text = "[Absent Frame]"
            row.cells[1].text = right_paras[j1] if j1 is not None else ""
            row.cells[2].text = "Injected External Entry Block clause verified."
            
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# ==========================================
# BLOCK 7: UI PHASE 1 - PORTAL INGESTION
# ==========================================
def render_due_diligence_landing():
    st.markdown('<div class="landing-wrapper">', unsafe_allow_html=True)
    st.markdown('<div class="brand-title">DELTA DUE DILIGENCE</div>', unsafe_allow_html=True)
    st.markdown('<div class="brand-subtitle">Asset Land Title Infrastructure Pipeline</div>', unsafe_allow_html=True)
    
    uploaded_files = st.file_uploader("Ingest Real Estate Assays, Title Documents or Title Photos", 
                                      type=['pdf', 'docx', 'png', 'jpg', 'jpeg', 'tiff'], 
                                      accept_multiple_files=True, label_visibility="collapsed")
    st.markdown('</div>', unsafe_allow_html=True)
    
    if uploaded_files:
        load_due_diligence_matrices(uploaded_files)
        st.markdown("<div style='max-width:600px; margin: 2rem auto 0 auto;'>", unsafe_allow_html=True)
        for filename in st.session_state.title_order:
            col1, col2 = st.columns([6, 4])
            with col1: 
                st.markdown(f'<p style="font-size:13px; color:#1A1A1A; padding-top:8px;">📁 {filename}</p>', unsafe_allow_html=True)
            with col2:
                roles = ["Certified True Copy (Base)", "Owner Duplicate (Counter)", "Registry Secondary Copy", "Assessed Tax Declaration Template"]
                st.session_state.title_roles[filename] = st.selectbox(f"Type_{filename}", roles, index=0, label_visibility="collapsed", key=f"dd_role_{filename}")
        st.markdown("<br/>", unsafe_allow_html=True)
        if st.button("Initialize Deep Variance Audit"):
            st.session_state.dd_processing_complete = True
            st.rerun()
        st.markdown("</div>", unsafe_allow_html=True)

# ==========================================
# BLOCK 8: UI PHASE 2 - SYSTEM DASHBOARD REVIEW
# ==========================================
def render_due_diligence_workspace():
    st.markdown('<div class="brand-title">DELTA DUE DILIGENCE WORKSPACE</div>', unsafe_allow_html=True)
    st.markdown("<br/>", unsafe_allow_html=True)
    
    ordered_files = st.session_state.title_order
    roles = [st.session_state.title_roles[f] for f in ordered_files]
    
    tab_review, tab_geospatial = st.tabs(["❖ Core Title Discrepancy Matrix", "⚡ Boundary Technical Description Plotter"])
    
    with tab_review:
        t_col1, t_col2 = st.columns(2)
        with t_col1: 
            st.session_state.base_title_idx = st.selectbox("Select Baseline Framework Core", range(len(ordered_files)), format_func=lambda x: f"BASE // {ordered_files[x]} ({roles[x]})")
        with t_col2: 
            st.session_state.counter_title_idx = st.selectbox("Select Target Evaluation Framework", range(len(ordered_files)), format_func=lambda x: f"EVAL // {ordered_files[x]} ({roles[x]})", index=min(1, len(ordered_files)-1))
            
        base_file = ordered_files[st.session_state.base_title_idx]
        counter_file = ordered_files[st.session_state.counter_title_idx]
        
        left_paras = st.session_state.uploaded_titles[base_file]
        right_paras = st.session_state.uploaded_titles[counter_file]
        
        alignment_opcodes = compute_fuzzy_alignment_matrix(left_paras, right_paras)
        
        st.markdown(f"""
            <div class="crypto-banner">
                INTEGRITY REGISTRY CRYPTO STAMP IDENTITY VALIDATION LOG<br/>
                [BASE_SHA256]: {st.session_state.title_hashes.get(base_file, "N/A")}<br/>
                [EVAL_SHA256]: {st.session_state.title_hashes.get(counter_file, "N/A")}
            </div>
        """, unsafe_allow_html=True)
        
        # FIXED: Expanded layouts allocations to clear text truncation bugs
        h_col1, _, h_col2, _ = st.columns([5, 0.2, 5, 4.5])
        with h_col1: st.markdown("<p style='font-size:11px; text-transform:uppercase; color:#737373; font-weight:600; margin-bottom:15px;'>BASELINE CORE SOURCE TEXT</p>", unsafe_allow_html=True)
        with h_col2: st.markdown("<p style='font-size:11px; text-transform:uppercase; color:#737373; font-weight:600; margin-bottom:15px;'>EVALUATION COMPARISON FIELD TARGET</p>", unsafe_allow_html=True)
        
        change_index = 1
        for idx, (tag, i1, _, j1, _) in enumerate(alignment_opcodes):
            # FIXED: Wider text panes allocation, streamlined review panels boundaries
            r_col1, r_col2, r_col3, r_col4 = st.columns([5, 0.2, 5, 4.5])
            
            with r_col1:
                if tag == 'equal':
                    st.markdown(f'<div class="doc-cell"><p class="stream-paragraph">{left_paras[i1]}</p></div>', unsafe_allow_html=True)
                elif tag == 'replace':
                    h1, _ = compute_token_diff_html(left_paras[i1], right_paras[j1])
                    st.markdown(f'<div class="doc-cell"><span class="trace-flag">▲ Variance Identified</span><p class="stream-paragraph">{h1}</p></div>', unsafe_allow_html=True)
                elif tag == 'delete':
                    st.markdown(f'<div class="doc-cell"><span class="trace-flag" style="color:#991B1B; border-bottom-color:#FEE2E2;">◼ Missing Element</span><p class="stream-paragraph"><span class="del-token">{left_paras[i1]}</span></p></div>', unsafe_allow_html=True)
                elif tag == 'insert':
                    st.markdown('<div class="doc-cell empty-cell"></div>', unsafe_allow_html=True)
                    
            with r_col2:
                if tag == 'equal': st.markdown('<div class="minimap-row minimap-equal"></div>', unsafe_allow_html=True)
                elif tag == 'replace': st.markdown('<div class="minimap-row minimap-replace"></div>', unsafe_allow_html=True)
                elif tag == 'delete': st.markdown('<div class="minimap-row minimap-delete"></div>', unsafe_allow_html=True)
                elif tag == 'insert': st.markdown('<div class="minimap-row minimap-insert"></div>', unsafe_allow_html=True)
                
            with r_col3:
                if tag == 'equal':
                    st.markdown(f'<div class="doc-cell"><p class="stream-paragraph">{right_paras[j1]}</p></div>', unsafe_allow_html=True)
                elif tag == 'replace':
                    _, h2 = compute_token_diff_html(left_paras[i1], right_paras[j1])
                    st.markdown(f'<div class="doc-cell"><span class="trace-flag">▲ Variance Identified</span><p class="stream-paragraph">{h2}</p></div>', unsafe_allow_html=True)
                elif tag == 'delete':
                    st.markdown('<div class="doc-cell empty-cell"></div>', unsafe_allow_html=True)
                elif tag == 'insert':
                    st.markdown(f'<div class="doc-cell"><span class="trace-flag" style="color:#065F46; border-bottom-color:#D1FAE5;">◆ External Injection</span><p class="stream-paragraph"><span class="add-token">{right_paras[j1]}</span></p></div>', unsafe_allow_html=True)
                    
            with r_col4:
                if tag != 'equal':
                    unique_id = f"dd_row_{i1}_{j1}_{idx}"
                    st.markdown(f'<div class="advisory-panel" style="margin-bottom:0px; padding-bottom:8px;"><div class="advisory-header"><span style="color:#EF4444;">🚨 Operational Alert #{change_index}</span></div><p style="font-size:11px; margin:0 0 8px 0; color:#525252;">Verify values for consistency.</p></div>', unsafe_allow_html=True)
                    st.radio("Risk Disposition", ["Clear", "Minor Variance", "Fraud Risk", "Hold Asset"], key=f"risk_{unique_id}", horizontal=True, label_visibility="collapsed")
                    st.text_input("Discrepancy Note Ledger", key=f"note_dd_{unique_id}", placeholder="Add due diligence findings...", label_visibility="collapsed")
                    change_index += 1
                else:
                    st.write("")
                    
        st.markdown("<br/>", unsafe_allow_html=True)
        export_bytes = export_due_diligence_docx(left_paras, right_paras, base_file, counter_file, alignment_opcodes)
        st.download_button("📥 Export Comprehensive Land Title Audit Matrix (.docx)", data=export_bytes, file_name=f"DELTA_DD_Matrix_{datetime.now().strftime('%Y%m%d')}.docx")

    with tab_geospatial:
        st.markdown("### Boundary Survey Line Geometry Transform Engine Matrix")
        st.markdown("Paste or modify the raw technical description text string segment block isolated from the land title below to evaluate mathematical polygon area parameters.")
        
        default_tech_desc = """THENCE S. 32 DEG. 00'E., 18.58 M. TO POINT 2; THENCE S. 69 DEG. 49'W., 50.64 M. TO POINT 3; THENCE N. 5 DEG. 19'W., 10.19 M. TO POINT 4; THENCE N. 35 DEG. 56'W., 7.77 M. TO POINT 5; THENCE N. 35 DEG. 56'W., 1.96 M. TO POINT 6; THENCE N. 28 DEG. 59'W., 9.72 M. TO POINT 7; THENCE N. 82 DEG. 12'E., 49.55 M. TO THE POINT OF BEGINNING;"""
        
        input_desc = st.text_area("Technical Description Script Parser Array Window Input Frame", value=default_tech_desc, height=180)
        
        if input_desc:
            coords = parse_technical_description(input_desc)
            if len(coords) > 1:
                st.success(f"Successfully tracked and vector structured [{len(coords)-1}] distinct vertex vectors boundaries nodes coordinates points loops arrays segments.")
                
                geo_col1, geo_col2 = st.columns([6, 4])
                
                with geo_col1:
                    fig, ax = plt.subplots(figsize=(6, 6), facecolor='#FFFFFF')
                    x_pts = [c[0] for c in coords]
                    y_pts = [c[1] for c in coords]
                    
                    ax.plot(x_pts, y_pts, color='#1A1A1A', linestyle='-', linewidth=1.5, marker='o', markerfacecolor='#FFFFFF', markeredgecolor='#1A1A1A', markersize=5)
                    for i, (xp, yp) in enumerate(coords[:-1]):
                        ax.text(xp + 1, yp + 1, f"P{i+1}", fontsize=8, fontweight='bold')
                        
                    ax.set_aspect('equal', 'box')
                    ax.grid(True, color='#E5E7EB', linestyle='--', linewidth=0.5)
                    ax.set_facecolor('#FFFFFF')
                    ax.spines['top'].set_visible(False)
                    ax.spines['right'].set_visible(False)
                    ax.spines['left'].set_color('#CCCCCC')
                    ax.spines['bottom'].set_color('#CCCCCC')
                    plt.title("TRANSFORM PLOT RADIAL AZIMUTH COMPASS PLAN PREVIEW", fontsize=10, fontname="DejaVu Sans", fontweight='bold', pad=12)
                    st.pyplot(fig)
                    plt.close(fig)
                    
                with geo_col2:
                    st.markdown("#### Export Pipeline Asset Formats Matrix")
                    st.markdown("Select an automated execution routine format pipeline down below to generate legal blueprint assets for CAD, GIS, or Title Presentation verification portfolios.")
                    
                    pdf_bytes = generate_white_blueprint_pdf(coords)
                    st.download_button(label="📥 Export Lot Plan White Plain PDF Document Asset",
                                       data=pdf_bytes,
                                       file_name=f"DELTA_LOT_PLAN_{datetime.now().strftime('%Y%m%d')}.pdf",
                                       mime="application/pdf")
                    
                    st.markdown("---")
                    
                    kml_bytes = generate_kml_payload(coords)
                    st.download_button(label="📥 Export Geospatial Vector KML (Google Earth/QGIS Engine)",
                                       data=kml_bytes,
                                       file_name=f"DELTA_LOT_GEOSPATIAL_{datetime.now().strftime('%Y%m%d')}.kml",
                                       mime="application/vnd.google-earth.kml+xml")
                    
                    st.markdown("---")
                    
                    st.markdown("**Local Relational Cartesian Bounds Registry Table Array Map**")
                    st.dataframe([{"Vertex Sequence Point": f"Point {idx+1}", "Relative Delta X (Eastings M.)": f"{pt[0]:.2f} m", "Relative Delta Y (Northings M.)": f"{pt[1]:.2f} m"} for idx, pt in enumerate(coords[:-1])], use_container_width=True)
            else:
                st.error("Engine failed to resolve bearing metrics tokens inside input description text. Check system punctuation bounds layout arrays formatting definitions.")

    st.markdown("<br/><br/><hr style='border-color:#E0E0E0;'/>", unsafe_allow_html=True)
    if st.button("Terminate Due Diligence Work Staging Context Array Frame Session"):
        st.session_state.dd_processing_complete = False
        st.rerun()
        
# ==========================================
# BLOCK 9: ENTRYPOINT ORCHESTRATION
# ==========================================
def main():
    st.set_page_config(page_title="DELTA DUE DILIGENCE PIPELINE", layout="wide", initial_sidebar_state="collapsed")
    inject_luxury_system_css()
    if not st.session_state.dd_processing_complete: 
        render_due_diligence_landing()
    else: 
        render_due_diligence_workspace()

if __name__ == "__main__": 
    main()
